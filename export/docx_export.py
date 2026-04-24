import os
import logging
from io import BytesIO
from datetime import datetime

import markdown as md_lib
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup, NavigableString, Tag
from PIL import Image

logger = logging.getLogger(__name__)

_BLUE = RGBColor(0x1F, 0x6F, 0xEB)
_BLUE_HEX = "1F6FEB"
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_GREY = RGBColor(0x66, 0x66, 0x66)
_DARK = RGBColor(0x1A, 0x1A, 0x2E)
_HEADING_SIZES = {"h1": 16, "h2": 15, "h3": 13, "h4": 11}


def _add_inline(para, element, bold=False, italic=False, code=False):
    if isinstance(element, NavigableString):
        text = str(element)
        if text:
            run = para.add_run(text)
            run.font.size = Pt(10)
            run.font.color.rgb = _DARK
            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True
            if code:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
    elif isinstance(element, Tag):
        if element.name == "br":
            para.add_run("\n")
            return
        # Handle internal anchor links (TOC entries)
        if element.name == "a":
            href = element.get("href", "")
            if href.startswith("#"):
                bookmark_name = href[1:]
                link_text = element.get_text(strip=True)
                if bookmark_name and link_text:
                    _add_hyperlink(para, bookmark_name, link_text)
                    return
            # Skip <a id="..."> anchor tags (bookmarks, no visible text needed)
            if element.get("id") and not element.get_text(strip=True):
                return
        nb = bold or element.name in ("strong", "b")
        ni = italic or element.name in ("em", "i")
        nc = code or element.name == "code"
        for child in element.children:
            _add_inline(para, child, bold=nb, italic=ni, code=nc)


def _set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    tcPr.append(shd)


def _set_table_full_width_centered(table):
    """Force a table to span 100% of the content width, autofit columns, and centre-align."""
    tbl = table._tbl
    tblPr = tbl.get_or_add_tblPr()

    # Remove any pre-existing tblW / jc / tblLayout elements to avoid duplicates
    for tag_name in ("w:tblW", "w:jc", "w:tblLayout"):
        for existing in tblPr.findall(qn(tag_name)):
            tblPr.remove(existing)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")   # 5000/100 = 100% in Word's pct unit
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    # autofit: Word distributes column widths based on content within the table width
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "autofit")
    tblPr.append(tblLayout)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)


def _add_para_border(para, side: str, color_hex: str, sz: str = "6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr = OxmlElement(f"w:{side}")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), sz)
    bdr.set(qn("w:space"), "4")
    bdr.set(qn("w:color"), color_hex)
    pBdr.append(bdr)
    pPr.append(pBdr)


def _slugify(text: str) -> str:
    """Convert heading text to a URL-friendly slug for bookmarks."""
    import re
    text = re.sub(r'<[^>]+>', '', text)
    text = text.lower().strip()
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[\s]+', '-', text)
    text = re.sub(r'-+', '-', text)
    return text.strip('-')


def _add_bookmark(para, bookmark_name: str):
    """Add a bookmark to a paragraph so it can be linked from the TOC."""
    tag = para._p
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(id(para) % 100000))
    bookmark_start.set(qn("w:name"), bookmark_name)
    tag.insert(0, bookmark_start)
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(id(para) % 100000))
    tag.append(bookmark_end)


def _add_hyperlink(para, bookmark_name: str, text: str):
    """Add an internal hyperlink (to a bookmark) in a paragraph."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _BLUE_HEX)
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "none")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")  # 10pt
    rPr.append(sz)
    run_el.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run_el.append(t)
    hyperlink.append(run_el)
    para._p.append(hyperlink)


def _load_image_as_png(path: str) -> BytesIO | None:
    """Load any image format and convert to PNG buffer."""
    if not path or not os.path.exists(path):
        return None
    try:
        img = Image.open(path).convert("RGBA")
        buf = BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf
    except Exception as exc:
        logger.warning(f"Image load failed for {path}: {exc}")
        return None


def _get_report_month(report_date: str) -> str:
    try:
        parts = report_date.split(" to ")
        if parts:
            dt = datetime.strptime(parts[-1].strip(), "%Y-%m-%d")
            return dt.strftime("%B %Y")
    except Exception:
        pass
    return report_date


def _add_cover_page(doc, customer_name: str, report_date: str,
                    logo_path: str | None, logicalis_logo: str):
    """Add a professional cover page."""
    report_month = _get_report_month(report_date)

    # Add some spacing at top
    for _ in range(4):
        doc.add_paragraph()

    # Logicalis logo
    logo_buf = _load_image_as_png(logicalis_logo)
    if logo_buf:
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.add_run().add_picture(logo_buf, height=Pt(48))

    # Customer logo if available
    if logo_path:
        cust_buf = _load_image_as_png(logo_path)
        if cust_buf:
            cust_para = doc.add_paragraph()
            cust_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cust_para.add_run().add_picture(cust_buf, height=Pt(48))

    doc.add_paragraph()

    # "Prepared by Logicalis for"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared by Logicalis for")
    run.font.size = Pt(14)
    run.font.color.rgb = _GREY

    # Customer name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(customer_name)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = _DARK

    doc.add_paragraph()

    # Divider
    div_para = doc.add_paragraph()
    div_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para_border(div_para, "bottom", _BLUE_HEX, sz="12")

    doc.add_paragraph()

    # "Logicalis Managed Security Services"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Logicalis Managed Security Services")
    run.font.size = Pt(14)
    run.font.color.rgb = _GREY

    # Report title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"GSOC Monthly Report \u2013 {report_month}")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = _BLUE

    doc.add_paragraph()

    # Page break after cover
    doc.add_page_break()


def _add_header_footer(doc, report_date: str):
    """Add header and footer to the document sections."""
    report_month = _get_report_month(report_date)

    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = ""
        run = header_para.add_run(f"Logicalis GSOC Monthly Report \u2013 {report_month}")
        run.font.size = Pt(8)
        run.font.color.rgb = _GREY
        _add_para_border(header_para, "bottom", "CCCCCC", sz="4")

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = ""
        run = footer_para.add_run("Confidential")
        run.font.size = Pt(8)
        run.font.color.rgb = _GREY

        # Add page number to footer right-aligned via tab
        run2 = footer_para.add_run("\t\t")
        run2.font.size = Pt(8)

        # Add page number field
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        run3 = footer_para.add_run()
        run3._r.append(fld_char1)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run4 = footer_para.add_run()
        run4._r.append(instr)
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run5 = footer_para.add_run()
        run5._r.append(fld_char2)


# Map chart names to section heading keywords for injection
_CHART_SECTION_MAP = {
    "monthly_trend": "1.2",
    "severity": "1.3",
    "resolution": "1.4",
    "sentinel_utilization": "1.8",
    "top_alerts": "1.9",
}


def generate_docx(markdown_content: str, customer_name: str, report_date: str,
                  logo_path: str | None = None, charts: dict | None = None) -> bytes:
    doc = Document()

    # Clear default paragraphs
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Paths
    default_logo = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "static", "Logo.webp"
    )

    # Add cover page
    _add_cover_page(doc, customer_name, report_date, logo_path, default_logo)

    # Add headers and footers
    _add_header_footer(doc, report_date)

    # Parse markdown to HTML
    content_html = md_lib.markdown(
        markdown_content,
        extensions=["tables", "fenced_code", "nl2br"],
    )
    soup = BeautifulSoup(content_html, "html.parser")

    # Build chart injection map: section_id -> list of chart PNG bytes
    chart_inject = {}
    if charts:
        for chart_name, png_bytes in charts.items():
            if not png_bytes:
                continue
            section_id = _CHART_SECTION_MAP.get(chart_name)
            if section_id:
                chart_inject.setdefault(section_id, []).append(png_bytes)

    def _maybe_insert_charts(heading_text: str):
        """Insert any charts that match this heading's section number."""
        if not chart_inject:
            return
        for section_id, chart_list in list(chart_inject.items()):
            if section_id in heading_text:
                for png_bytes in chart_list:
                    try:
                        img_buf = BytesIO(png_bytes)
                        chart_para = doc.add_paragraph()
                        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        chart_para.add_run().add_picture(img_buf, width=Inches(5.5))
                    except Exception as exc:
                        logger.warning(f"Chart embed failed: {exc}")
                del chart_inject[section_id]

    # Process each HTML element
    for element in soup.children:
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                doc.add_paragraph(text)
            continue
        if not isinstance(element, Tag):
            continue

        name = element.name

        if name in _HEADING_SIZES:
            level = int(name[1])
            heading_text = element.get_text(strip=True)
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
            para.paragraph_format.space_after = Pt(6 if level <= 2 else 4)
            run = para.add_run(heading_text)
            run.font.size = Pt(_HEADING_SIZES[name])
            run.font.bold = True
            run.font.color.rgb = _BLUE
            # Add bottom border for h2
            if level == 2:
                _add_para_border(para, "bottom", _BLUE_HEX, sz="8")
            # Add bookmark for TOC linking
            bookmark_id = _slugify(heading_text)
            # Also check for <a id="..."> inside the heading
            anchor_tag = element.find("a", attrs={"id": True})
            if anchor_tag:
                bookmark_id = anchor_tag["id"]
            if bookmark_id:
                _add_bookmark(para, bookmark_id)
            # Insert charts after matching headings
            _maybe_insert_charts(heading_text)

        elif name == "p":
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(4)
            _add_inline(para, element)

        elif name in ("ul", "ol"):
            for i, li in enumerate(element.find_all("li", recursive=False)):
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.3)
                para.paragraph_format.space_after = Pt(2)
                bullet = f"{i + 1}." if name == "ol" else "\u2022"
                run = para.add_run(f"{bullet} ")
                run.font.size = Pt(10)
                _add_inline(para, li)

        elif name == "pre":
            code_el = element.find("code")
            code_text = code_el.get_text() if code_el else element.get_text()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.2)
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F0F3FA")
            shd.set(qn("w:val"), "clear")
            pPr.append(shd)
            run = para.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        elif name == "table":
            rows_data = []
            for tr in element.find_all("tr"):
                cells = tr.find_all(["th", "td"])
                rows_data.append([(c.get_text(strip=True), c.name == "th") for c in cells])
            if not rows_data:
                continue
            ncols = max(len(r) for r in rows_data)
            table = doc.add_table(rows=len(rows_data), cols=ncols)
            table.style = "Table Grid"

            for ri, row_data in enumerate(rows_data):
                for ci, (text, is_th) in enumerate(row_data):
                    if ci >= ncols:
                        continue
                    cell = table.cell(ri, ci)
                    cell_para = cell.paragraphs[0]
                    run = cell_para.add_run(text)
                    run.font.size = Pt(8)

                    if is_th or ri == 0:
                        # Header row: blue background, white text
                        run.font.bold = True
                        run.font.color.rgb = _WHITE
                        _set_cell_shading(cell, _BLUE_HEX)
                    elif ri % 2 == 0:
                        # Alternating rows
                        _set_cell_shading(cell, "F7F9FF")

            _set_table_full_width_centered(table)
            doc.add_paragraph()

        elif name == "blockquote":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            # Amber/orange left border for placeholder blocks
            _add_para_border(para, "left", "F59E0B", sz="16")
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "FFFBEB")
            shd.set(qn("w:val"), "clear")
            pPr.append(shd)
            _add_inline(para, element)

        elif name == "hr":
            para = doc.add_paragraph()
            _add_para_border(para, "bottom", "DDDDDD", sz="4")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()
